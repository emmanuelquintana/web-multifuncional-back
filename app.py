from flask import Flask, request, jsonify, send_from_directory
from werkzeug.utils import secure_filename
import os
from flask_cors import CORS
import fitz
from docx import Document
from docx.shared import Cm
from PIL import Image
from datetime import datetime
from docx2pdf import convert
import math
import shutil
import threading
import time
import json

app = Flask(__name__)
CORS(app)

# Configurar la carpeta de subida
UPLOAD_FOLDER = 'uploads'
if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

# Configurar la ruta de descarga
DOWNLOAD_FOLDER = 'download'
if not os.path.exists(DOWNLOAD_FOLDER):
    os.makedirs(DOWNLOAD_FOLDER)
app.config['DOWNLOAD_FOLDER'] = DOWNLOAD_FOLDER


# Configurar el registro (logs)
import logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# Función para convertir PDF a imágenes JPG
def pdf_to_jpg(pdf_path, output_folder, min_width, min_height, day):
    
    pdf_document = fitz.open(pdf_path)
    total_pages = len(pdf_document)
    pedidos = total_pages // 2

    for page_number in range(total_pages):
        page = pdf_document.load_page(page_number)
        pix = page.get_pixmap()
        if pix.width < min_width or pix.height < min_height:
            scale_factor = max(min_width / pix.width, min_height / pix.height)
            pix = page.get_pixmap(matrix=fitz.Matrix(scale_factor, scale_factor))
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
        output_path = os.path.join(output_folder, f"{page_number + 1:04d}.jpg")
        img.save(output_path, quality=100)
        logger.info(f"Guardado {output_path}")

    pdf_document.close()
    return total_pages, pedidos

# Función para procesar imágenes en un documento Word y convertirlo a PDF
def process_images(image_folder, day):
    doc = Document()
    for section in doc.sections:
        section.left_margin = Cm(0.5)
        section.right_margin = Cm(0.5)
        section.top_margin = Cm(0.5)
        section.bottom_margin = Cm(0.5)

    image_files = sorted([f for f in os.listdir(image_folder) if f.endswith('.jpg')])
    desired_width = Cm(7.59)
    desired_height = Cm(13.02)
    image_count = 0
    total_imagenes_inicio = len(image_files)

    for image_file in image_files:
        img_path = os.path.join(image_folder, image_file)
        img = Image.open(img_path)
        img.thumbnail((desired_width, desired_height))
        if image_count == 0:
            table = doc.add_table(rows=2, cols=2)
            table.autofit = False
            table.allow_autofit = False
        cell = table.cell(image_count // 2, image_count % 2)
        run = cell.paragraphs[0].add_run()
        run.add_picture(img_path, width=desired_width, height=desired_height)
        image_count += 1
        if image_count == 4:
            doc.add_page_break()
            image_count = 0

    hojas_de_4_imagenes = math.ceil(total_imagenes_inicio / 4)
    fecha_actual = datetime.now().strftime("%d-%m-%Y")
    nombre_archivo_doc = f"Guias Shein {fecha_actual} {day} medidas pequeñas.docx"
    nombre_archivo_pdf = f"Guias Shein {fecha_actual} {day} medidas pequeñas.pdf"
    doc.save(nombre_archivo_doc)
    logger.info(f"El archivo Word '{nombre_archivo_doc}' ha sido creado con éxito.")
    convert(nombre_archivo_doc)
    logger.info(f"El archivo PDF '{nombre_archivo_pdf}' ha sido creado con éxito.")
    eliminar_hojas_pares(nombre_archivo_pdf)

    # Rutas relativas para descargar desde el front-end
    rel_path_doc = os.path.relpath(nombre_archivo_doc, os.getcwd())
    rel_path_pdf = os.path.relpath(nombre_archivo_pdf, os.getcwd())

    return total_imagenes_inicio // 2, rel_path_doc, rel_path_pdf

#Funcion Auxiliar para organizar Documentos en carpeta Download e ignorar archivos de la carpeta venv
def organizar_documentos():
    for root, dirs, files in os.walk(os.getcwd()):
        if 'venv' in root:
            continue  # Ignorar la carpeta venv
        if 'upload' in root:
            continue  # Ignorar la carpeta upload
        for file in files:
            if file.endswith('.pdf') or file.endswith('.docx'):
                shutil.move(os.path.join(root, file), os.path.join(app.config['DOWNLOAD_FOLDER'], file))


# Función para eliminar páginas pares de un PDF
def eliminar_hojas_pares(nombre_archivo_pdf):
    pdf_document = fitz.open(nombre_archivo_pdf)
    paginas_eliminadas = 0
    for index, _ in reversed(list(enumerate(pdf_document))):
        if (index + 1) % 2 == 0:
            pdf_document.delete_page(index)
            paginas_eliminadas += 1
    nombre_archivo_pdf_sin_pares = f"{nombre_archivo_pdf.split('.pdf')[0]}_sin_pares.pdf"
    pdf_document.save(nombre_archivo_pdf_sin_pares)
    pdf_document.close()
    logger.info(f"Se eliminaron {paginas_eliminadas} páginas pares del archivo '{nombre_archivo_pdf_sin_pares}'.")
    os.remove(nombre_archivo_pdf)
    logger.info("El PDF original ha sido eliminado.")
    #quitar _sin_pares del titulo del documento
    os.rename(nombre_archivo_pdf_sin_pares, nombre_archivo_pdf)
    
 # Función para renombrar y organizar la carpeta de subidas
def rename_and_organize_uploads():
    fecha_actual = datetime.now().strftime("%d-%m-%Y")
    for file in os.listdir(app.config['UPLOAD_FOLDER']):
        if file.endswith('.pdf'):
            old_path = os.path.join(app.config['UPLOAD_FOLDER'], file)
            new_filename = f"guias-{fecha_actual}-{file}"
            new_path = os.path.join(app.config['UPLOAD_FOLDER'], new_filename)
            os.rename(old_path, new_path)
            logger.info(f"Archivo renombrado a {new_filename}")

# Función para vaciar carpetas una vez a la semana
def clear_folders_weekly():
    while True:
        time.sleep(7 * 24 * 60 * 60)  # Esperar una semana
        clear_folders()
        logger.info("Carpetas 'uploads', 'download' y carpetas de imágenes vaciadas.")

# Función para vaciar las carpetas
def clear_folders():
    for folder in [app.config['UPLOAD_FOLDER'], app.config['DOWNLOAD_FOLDER']]:
        for file in os.listdir(folder):
            file_path = os.path.join(folder, file)
            try:
                if os.path.isfile(file_path) or os.path.islink(file_path):
                    os.unlink(file_path)
                elif os.path.isdir(file_path):
                    shutil.rmtree(file_path)
            except Exception as e:
                logger.error(f'Error al borrar el archivo {file_path}. Razón: {e}')

    # Eliminar carpetas de imágenes generadas
    current_dir = os.getcwd()
    for folder in os.listdir(current_dir):
        folder_path = os.path.join(current_dir, folder)
        if os.path.isdir(folder_path) and folder.startswith("GUIAS SHEIN") and folder.endswith("- IMAGENES"):
            try:
                shutil.rmtree(folder_path)
                logger.info(f'Carpeta {folder_path} eliminada.')
            except Exception as e:
                logger.error(f'Error al borrar la carpeta {folder_path}. Razón: {e}')

# Iniciar el hilo para vaciar las carpetas semanalmente
threading.Thread(target=clear_folders_weekly, daemon=True).start()


# Función para procesar un solo día
@app.route('/api/process_single_day', methods=['POST'])
def process_single_day():
    if 'file' not in request.files:
        return jsonify({'error': 'No se ha proporcionado ningún archivo'}), 400

    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'Nombre de archivo vacío'}), 400

    if file and file.filename.endswith('.pdf'):
        filename = secure_filename(file.filename)
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(file_path)

        output_folder_name = f"GUIAS SHEIN {datetime.now().strftime('%Y-%m-%d')} - IMAGENES"
        output_folder_path = os.path.join(os.getcwd(), output_folder_name)
        if not os.path.exists(output_folder_path):
            os.makedirs(output_folder_path)

        min_width = 896
        min_height = 1538
        total_pages, pedidos = pdf_to_jpg(file_path, output_folder_path, min_width, min_height, datetime.now().strftime("%d-%m-%Y"))
        total_pedidos = total_pages // 2


        processed_pedidos, rel_path_doc, rel_path_pdf = process_images(output_folder_path, datetime.now().strftime("%d-%m-%Y"))
        
          # Llamada a la función organizar_documentos()
        organizar_documentos()

        rename_and_organize_uploads()

        rel_path_doc_encoded = rel_path_doc.replace(' ', '%20')
        rel_path_pdf_encoded = rel_path_pdf.replace(' ', '%20')

        doc_download_link = f"/api/download/{rel_path_doc_encoded}"
        pdf_download_link = f"/api/download/{rel_path_pdf_encoded}"


        response = {
            'message': f'Procesado exitosamente para {datetime.now().strftime("%d-%m-%Y")}',
            'total_pages': total_pages,
            'total_pedidos': total_pedidos,
            'processed_pedidos': processed_pedidos,
            'name_doc': rel_path_doc,
            'name_pdf': rel_path_pdf,
            'links': {
                'download_doc': doc_download_link,
                'download_pdf': pdf_download_link
            }
        }

        logger.info(f"Respuesta JSON: {response}")

        return jsonify(response), 200

    return jsonify({'error': 'Formato de archivo no válido, se esperaba PDF'}), 400

# Función para procesar el fin de semana
@app.route('/api/process_weekend', methods=['POST'])
def process_weekend():
    if 'files[]' not in request.files:
        return jsonify({'error': 'No se han proporcionado archivos'}), 400

    files = request.files.getlist('files[]')
    if len(files) != 3:
        return jsonify({'error': 'Debe seleccionar tres archivos PDF (Viernes, Sábado y Domingo)'}), 400

    for file in files:
        if file.filename == '':
            return jsonify({'error': 'Nombre de archivo vacío'}), 400
        if not file.filename.endswith('.pdf'):
            return jsonify({'error': f'Archivo {file.filename} no es un PDF válido'}), 400

    day_files = {}
    for file in files:
        filename = secure_filename(file.filename)
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(file_path)
        day = get_day_from_filename(filename)
        day_files[day] = file_path

    processed_days = []
    for day, file_path in day_files.items():
        output_folder_name = f"GUIAS SHEIN {datetime.now().strftime('%Y-%m-%d')} - {day} - IMAGENES"
        output_folder_path = os.path.join(os.getcwd(), output_folder_name)
        if not os.path.exists(output_folder_path):
            os.makedirs(output_folder_path)

        min_width = 896
        min_height = 1538
        total_pages, pedidos = pdf_to_jpg(file_path, output_folder_path, min_width, min_height, day)
        total_pedidos = total_pages // 2

        processed_pedidos, rel_path_doc, rel_path_pdf = process_images(output_folder_path, day)

        # Llamada a la función organizar_documentos()
        organizar_documentos()

        #Solo hacer el replace si el archivo contiene espacios si no solo dejarlo como esta con un if
        if ' ' in rel_path_doc:
            rel_path_doc = rel_path_doc.replace(' ', '%20')
        if ' ' in rel_path_pdf:
            rel_path_pdf = rel_path_pdf.replace(' ', '%20')

        doc_download_link = f"/api/download/{rel_path_doc}"
        pdf_download_link = f"/api/download/{rel_path_pdf}"

        processed_days.append({
            'day': day,
            'total_pages': total_pages,
            'total_pedidos': total_pedidos,
            'processed_pedidos': processed_pedidos,
            'name_doc': rel_path_doc,
            'name_pdf': rel_path_pdf,
            'links': {
                'download_doc': doc_download_link,
                'download_pdf': pdf_download_link
            }
        })

    rename_and_organize_uploads()
    response = {
        'message': 'Procesado exitosamente para el fin de semana',
        'processed_days': processed_days
    }

    logger.info(f"Respuesta JSON: {json.dumps(response, indent=2)}")

    return jsonify(response), 200

# Función auxiliar para obtener el día desde el nombre del archivo
def get_day_from_filename(filename):
    if 'viernes' in filename.lower():
        return 'Viernes'
    elif 'sabado' in filename.lower() or 'sábado' in filename.lower():
        return 'Sábado'
    elif 'domingo' in filename.lower():
        return 'Domingo'
    else:
        return ''

# Ruta para descargar archivos generados# Ruta para descargar archivos generadosfrom flask import send_from_directory

# Función para procesar la descarga de archivos
@app.route('/api/download/<path:filename>')
def download_file(filename):
    directory = app.config['DOWNLOAD_FOLDER']
    logger.info("directory: " + directory)

    filename = filename.replace('%20', ' ')  # Reemplazar %20 por espacios en blanco
    file_path = os.path.join(directory, filename)  # Obtener la ruta completa del archivo
    
    if os.path.exists(file_path):
        return send_from_directory(directory=directory, path=filename, as_attachment=True)
    else:
        return jsonify({'error': f'Archivo {filename} no encontrado'}), 404

# Endpoint para vaciar las carpetas
@app.route('/api/clear_folders', methods=['POST'])
def clear_folders_endpoint():
    clear_folders()
    logger.info("Carpetas 'uploads' y 'download' vaciadas a través del endpoint.")
    return jsonify({'message': 'Carpetas vaciadas exitosamente'}), 200

if __name__ == '__main__':
    app.run(debug=True)
